import os
import sys
import json
import docx

# Reconfigure stdout and stderr for UTF-8 encoding on Windows console
if sys.platform.startswith('win'):
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')

from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH



def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._element = p._p = None

def format_paragraph_to_standard(p):
    # Set first line indent to 1cm (0.39 inches)
    p.paragraph_format.first_line_indent = Cm(1.0)
    # Format all runs to Times New Roman, Size 14
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def format_agency_list(items):
    """
    Format a list of agencies into standard administrative format:
    - Each item starts with "- "
    - Separated by ";"
    - The last item ends with "."
    - If only 1 item, no "-" at start.
    """
    if not items:
        return ""
    if isinstance(items, str):
        items = [i.strip() for i in items.replace('\n', ';').split(';') if i.strip()]
        
    clean_items = []
    for item in items:
        clean_item = item.strip().lstrip('-').strip().rstrip(';. ')
        if clean_item:
            clean_items.append(clean_item)
            
    if not clean_items:
        return ""
        
    if len(clean_items) == 1:
        return f"{clean_items[0]}."
        
    formatted = []
    for idx, item in enumerate(clean_items):
        if idx == len(clean_items) - 1:
            formatted.append(f"- {item}.")
        else:
            formatted.append(f"- {item};")
    return "\n".join(formatted)


def normalize_agency_name(name):
    """Normalize agency name to check for similarity."""
    if not name:
        return ""
    val = name.strip().lower().replace("  ", " ")
    val = val.replace("uỷ", "ủy").replace("oà", "òa").replace("uý", "úy")
    return val

def sort_agencies(agencies):
    def get_sort_key(name):
        norm = normalize_agency_name(name)
        if "ủy ban nhân dân" in norm or "ubnd" in norm:
            return 1
        elif "mặt trận" in norm or "mttq" in norm:
            return 2
        elif "xây dựng đảng" in norm or "xd đảng" in norm:
            return 3
        elif "kiểm tra" in norm or "ubkt" in norm:
            return 4
        return 5
    return sorted(agencies, key=get_sort_key)

def convert_to_old_tone_style(text):
    if not text:
        return text
    replacements = {
        "òa": "oà", "óa": "oá", "ỏa": "oả", "õa": "oã", "ọa": "oạ",
        "òe": "oè", "óe": "oé", "ỏe": "oẻ", "õe": "oẽ", "ọe": "oẹ",
        "úy": "uý", "ùy": "uỳ", "ủy": "uỷ", "ũy": "uỹ", "ụy": "uỵ",
        "Òa": "Oà", "Óa": "Oá", "Ỏa": "Oả", "Õa": "Oã", "Ọa": "Oạ",
        "Òe": "Oè", "Óe": "Oé", "Ỏe": "Oẻ", "Õe": "Oẽ", "Ọe": "Oẹ",
        "Úy": "Uý", "Ùy": "Uỳ", "Ủy": "Uỷ", "Ũy": "Uỹ", "Ụy": "Uỵ",
        "ÒA": "OÀ", "ÓA": "OÁ", "ỎA": "OẢ", "ÕA": "OÃ", "ỌA": "OẠ",
        "ÒE": "OÈ", "ÓE": "OÉ", "ỎE": "OẺ", "ÕE": "OẼ", "ỌE": "OẸ",
        "ÚY": "UÝ", "ÙY": "UỲ", "ỦY": "UỶ", "ŨY": "UỸ", "ỤY": "UỴ"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text

def apply_old_tone_style_to_doc(doc):
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = convert_to_old_tone_style(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = convert_to_old_tone_style(run.text)

def parse_parent_document_date(van_ban_cap_tren):
    import re
    from datetime import datetime
    # Match ngày DD/MM/YYYY or DD-MM-YYYY
    m1 = re.search(r'ngày\s+(\d{1,2})[/-](\d{1,2})[/-](\d{4})', van_ban_cap_tren, re.IGNORECASE)
    if m1:
        day = int(m1.group(1))
        month = int(m1.group(2))
        year = int(m1.group(3))
        return datetime(year, month, day)
        
    # Match ngày DD tháng MM năm YYYY
    m2 = re.search(r'ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})', van_ban_cap_tren, re.IGNORECASE)
    if m2:
        day = int(m2.group(1))
        month = int(m2.group(2))
        year = int(m2.group(3))
        return datetime(year, month, day)
        
    return datetime.now()

def calculate_default_deadline(parent_date_str):
    from datetime import timedelta, datetime
    today = datetime.now()
    # Tính 10 ngày làm việc (bỏ qua Thứ 7 và Chủ nhật)
    working_days_added = 0
    deadline = today
    while working_days_added < 10:
        deadline += timedelta(days=1)
        # 0 = Monday ... 4 = Friday (ngày làm việc)
        if deadline.weekday() < 5:
            working_days_added += 1
    return deadline.strftime("%d/%m/%Y")

def apply_bold_italic_to_deadline(paragraph, date_str):
    text = paragraph.text
    target_phrases = [f"trước ngày {date_str}", f"trước {date_str}"]
    target_phrase = None
    for phrase in target_phrases:
        if phrase in text:
            target_phrase = phrase
            break
            
    if not target_phrase:
        return
        
    start_idx = text.find(target_phrase)
    end_idx = start_idx + len(target_phrase)
    
    before_text = text[:start_idx]
    match_text = text[start_idx:end_idx]
    after_text = text[end_idx:]
    
    paragraph.text = ""
    
    if before_text:
        run_before = paragraph.add_run(before_text)
        run_before.font.name = 'Times New Roman'
        run_before.font.size = Pt(14)
        rPr = run_before._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        
    run_match = paragraph.add_run(match_text)
    run_match.bold = True
    run_match.italic = True
    run_match.font.name = 'Times New Roman'
    run_match.font.size = Pt(14)
    rPr = run_match._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    
    if after_text:
        run_after = paragraph.add_run(after_text)
        run_after.font.name = 'Times New Roman'
        run_after.font.size = Pt(14)
        rPr = run_after._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def format_all_deadlines_bold_italic(doc, dates):
    for date_str in dates:
        if not date_str:
            continue
        for p in doc.paragraphs:
            apply_bold_italic_to_deadline(p, date_str)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        apply_bold_italic_to_deadline(p, date_str)

def generate_document(data, template_path, output_path):
    # Load document
    doc = docx.Document(template_path)
    
    # Extract inputs
    van_ban_cap_tren = data.get("van_ban_cap_tren", "")
    
    # Extract short version of parent document (e.g. "Kế hoạch số 84-KH/TU, ngày 15/6/2026")
    import re
    van_ban_cap_tren_rut_gon = van_ban_cap_tren
    parts = re.split(r'\s+của\s+', van_ban_cap_tren, maxsplit=1, flags=re.IGNORECASE)
    if parts and len(parts) > 0:
        van_ban_cap_tren_rut_gon = parts[0].strip()
        
    # Determine if it's an eligible parent document type (Kế hoạch, Nghị quyết, Chương trình hành động)
    is_eligible = False
    norm_title = van_ban_cap_tren.lower()
    for keyword in ["kế hoạch", "nghị quyết", "chương trình"]:
        if keyword in norm_title:
            is_eligible = True
            break
            
    ngay_den_han_1 = data.get("Ngay_den_han_1", "")
    if (not ngay_den_han_1 or ngay_den_han_1.strip() == "") and is_eligible:
        ngay_den_han_1 = calculate_default_deadline(van_ban_cap_tren)
        
    co_quan_2 = data.get("Co_quan_2", "")
    
    ngay_den_han_2 = data.get("ngay_den_han_2", "")
    if (not ngay_den_han_2 or ngay_den_han_2.strip() == "") and is_eligible:
        ngay_den_han_2 = calculate_default_deadline(van_ban_cap_tren)
        
    trich_yeu = data.get("TRICH_YEU_CONG_VAN", "")
    
    # Check if co_quan_2 is Ban xây dựng Đảng (case insensitive)
    is_co_quan_2_ban_xd_dang = normalize_agency_name(co_quan_2) in [
        normalize_agency_name("Ban xây dựng Đảng"),
        normalize_agency_name("Ban Xây dựng Đảng")
    ]
    
    # Rule: "Chỉ gửi cho đơn vị được giao nhiệm vụ"
    # Assigned units are "Ban Xây dựng Đảng" (always task 1), and co_quan_2 if not merged.
    assigned_units = ["Ban Xây dựng Đảng"]
    if co_quan_2 and not is_co_quan_2_ban_xd_dang:
        # Keep original capitalization of co_quan_2
        assigned_units.append(co_quan_2)
        
    # Sort assigned units according to user rules
    assigned_units = sort_agencies(assigned_units)
        
    # Format lists
    danh_sach_gui = format_agency_list(assigned_units)
    danh_sach_noi_nhan = format_agency_list(data.get("DANH_SACH_NOI_NHAN", []))

    
    # We will search and replace in paragraphs and table cells
    # We will process paragraphs in order to handle the deletion/merging rule
    
    paragraphs_to_delete = []
    
    # Process all main paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text
        
        # Check consolidation rule for Paragraph 2 and 3
        # In our inspection:
        # Paragraph 2 text: '1. Giao Ban Xây dựng Đảng tham mưu Ban Thường vụ Đảng uỷ quán triệt, tuyên truyền {{van_ban_cap_tren}}; trình Thường trực Đảng uỷ trước {{Ngay_den_han_1}}.'
        # Paragraph 3 text: '2. Giao {{Co_quan_2}} tham mưu Ban Thường vụ Đảng uỷ triển khai thực hiện {{van_ban_cap_tren}}, đảm bảo phù hợp với tình hình thực tiễn của địa phương; trình Thường trực Đảng uỷ trước {{ngay_den_han_2}}.'
        
        if "Thực hiện {{van_ban_cap_tren}}" in text:
            new_text = text.replace("{{van_ban_cap_tren}}", van_ban_cap_tren)
            p.text = new_text
            format_paragraph_to_standard(p)
            
        elif "1. Giao Ban Xây dựng Đảng" in text:
            if is_co_quan_2_ban_xd_dang:
                # Merge tasks (no number prefix when there is only 1 task)
                new_text = f"Giao Ban Xây dựng Đảng tham mưu Ban Thường vụ Đảng uỷ quán triệt, tuyên truyền và triển khai thực hiện {van_ban_cap_tren_rut_gon}; trình Thường trực Đảng uỷ trước {ngay_den_han_1}."
            else:
                new_text = text.replace("{{van_ban_cap_tren_rut_gon}}", van_ban_cap_tren_rut_gon)\
                               .replace("{{van_ban_cap_tren}}", van_ban_cap_tren_rut_gon)\
                               .replace("{{Ngay_den_han_1}}", ngay_den_han_1)
            p.text = new_text
            format_paragraph_to_standard(p)
            
        elif "2. Giao {{Co_quan_2}}" in text:
            if is_co_quan_2_ban_xd_dang:
                # Mark for deletion
                paragraphs_to_delete.append(p)
            else:
                new_text = text.replace("{{Co_quan_2}}", co_quan_2)\
                               .replace("{{van_ban_cap_tren_rut_gon}}", van_ban_cap_tren_rut_gon)\
                               .replace("{{van_ban_cap_tren}}", van_ban_cap_tren_rut_gon)\
                               .replace("{{ngay_den_han_2}}", ngay_den_han_2)
                p.text = new_text
                format_paragraph_to_standard(p)
        else:
            # General replacement
            replaced = False
            if "{{van_ban_cap_tren_rut_gon}}" in text:
                text = text.replace("{{van_ban_cap_tren_rut_gon}}", van_ban_cap_tren_rut_gon)
                replaced = True
            if "{{van_ban_cap_tren}}" in text:
                text = text.replace("{{van_ban_cap_tren}}", van_ban_cap_tren_rut_gon)
                replaced = True
            if "{{Ngay_den_han_1}}" in text:
                text = text.replace("{{Ngay_den_han_1}}", ngay_den_han_1)
                replaced = True
            if "{{Co_quan_2}}" in text:
                text = text.replace("{{Co_quan_2}}", co_quan_2)
                replaced = True
            if "{{ngay_den_han_2}}" in text:
                text = text.replace("{{ngay_den_han_2}}", ngay_den_han_2)
                replaced = True
            if replaced:
                p.text = text
                
    # Delete marked paragraphs
    for p in paragraphs_to_delete:
        delete_paragraph(p)
        
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text
                    replaced = False
                    
                    if "{{TRICH_YEU_CONG_VAN}}" in text:
                        text = text.replace("{{TRICH_YEU_CONG_VAN}}", trich_yeu)
                        p.text = text
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.font.italic = True
                            rPr = run._r.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            
                    elif "- {{DANH_SACH_GUI}};" in text or "{{DANH_SACH_GUI}}" in text:
                        text = text.replace("- {{DANH_SACH_GUI}};", danh_sach_gui).replace("{{DANH_SACH_GUI}}", danh_sach_gui)
                        p.text = text
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in p.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(14)
                            rPr = run._r.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                            
                    else:
                        if "{{DANH_SACH_NOI_NHAN}}" in text:
                            delete_paragraph(p)
                            continue

                            
                        # Standard fallback replacements inside tables
                        if "{{van_ban_cap_tren_rut_gon}}" in text:
                            text = text.replace("{{van_ban_cap_tren_rut_gon}}", van_ban_cap_tren_rut_gon)
                            replaced = True
                        if "{{van_ban_cap_tren}}" in text:
                            text = text.replace("{{van_ban_cap_tren}}", van_ban_cap_tren_rut_gon)
                            replaced = True
                        if "{{Ngay_den_han_1}}" in text:
                            text = text.replace("{{Ngay_den_han_1}}", ngay_den_han_1)
                            replaced = True
                        if "{{Co_quan_2}}" in text:
                            text = text.replace("{{Co_quan_2}}", co_quan_2)
                            replaced = True
                        if "{{ngay_den_han_2}}" in text:
                            text = text.replace("{{ngay_den_han_2}}", ngay_den_han_2)
                            replaced = True
                            
                        if replaced:
                            p.text = text

    # Format all deadlines to bold and italic: "trước ngày [ngày đến hạn]"
    format_all_deadlines_bold_italic(doc, [ngay_den_han_1, ngay_den_han_2])

    # Apply tone conversion to match administrative standard: oà, uý, uỷ
    apply_old_tone_style_to_doc(doc)

    # Save output
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    try:
        doc.save(output_path)
        print(f"SUCCESS: Document generated at {output_path}")
    except PermissionError:
        import time
        suffix = int(time.time())
        base, ext = os.path.splitext(output_path)
        fallback_path = f"{base}_{suffix}{ext}"
        doc.save(fallback_path)
        print(f"WARNING: File {output_path} is locked. Saved as fallback to {fallback_path}")
        print(f"SUCCESS: Document generated at {fallback_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        # Help or demo run if executed directly without args
        print("Usage: python generate_docx.py <data_json_path> <output_path>")
        
        # We can run a demo test to verify it works
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        demo_template = os.path.join(base_dir, "references", "cong_van_giao_viec_mau.docx")
        demo_output = os.path.join(base_dir, "output", "Demo_Cong_Van_Giao_Viec.docx")
        
        demo_data = {
            "van_ban_cap_tren": "Nghị quyết số 12-NQ/TW, ngày 15/06/2026 của Ban Chấp hành Trung ương về nâng cao hiệu quả tuyên truyền trong thời kỳ mới",
            "Ngay_den_han_1": "15/07/2026",
            "Co_quan_2": "Ban xây dựng Đảng", # Will trigger the merge rule
            "ngay_den_han_2": "20/07/2026",
            "TRICH_YEU_CONG_VAN": "V/v thực hiện Nghị quyết số 12-NQ/TW của Ban Chấp hành Trung ương",
            "DANH_SACH_GUI": ["Ban Xây dựng Đảng", "Ủy ban Kiểm tra Đảng ủy xã", "UBND xã", "Ủy ban Mặt trận Tổ quốc Việt Nam xã"],
            "DANH_SACH_NOI_NHAN": ["Thường trực Đảng ủy", "Các đồng chí Ủy viên BCH", "Ủy ban kiểm tra", "VP HĐND & UBND"]
        }
        
        print("\nRunning test generation with merge rule...")
        generate_document(demo_data, demo_template, demo_output)
    else:
        json_path = sys.argv[1]
        out_path = sys.argv[2]
        
        with open(json_path, 'r', encoding='utf-8') as f:
            input_data = json.load(f)
            
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_p = os.path.join(base_dir, "references", "cong_van_giao_viec_mau.docx")
        generate_document(input_data, template_p, out_path)
